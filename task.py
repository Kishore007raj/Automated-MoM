# Step 1: Importing the required packages
# Import essential libraries for document processing, summarization, and output formatting.

import os
from llama_index.core import SimpleDirectoryReader # To read documents from a directory.
from llama_index.core import SummaryIndex # To create and manage summaries.
from llama_index.llms.groq import Groq # Groq model for generating summaries.
from llama_index.core.node_parser import SentenceSplitter # For splitting text into manageable chunks.
from llama_index.core import Settings # To configure model settings.
from llama_index.embeddings.huggingface import HuggingFaceEmbedding # Embeddings for better representation.
from docx import Document  # To create and save Word documents.

# Step 2: Load the API key from the .env file
# Securely load API key from a hidden .env file using python-dotenv.

from dotenv import load_dotenv
load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
os.environ["GROQ_API_KEY"] = GROQ_API_KEY

# Step 3: Enable asynchronous functionality
# Apply nest_asyncio to handle asynchronous functions smoothly in the script.

import nest_asyncio
nest_asyncio.apply()

# step 4: Initialize the Groq model
# Set up the Groq model for summarization.

llm = Groq(model="llama-3.1-8b-instant")
Settings.llm = Groq(model="llama-3.1-8b-instant")
Settings.embed_model = HuggingFaceEmbedding()

# step 5: Load the documents
# Read the input document (PDF) for summarization.

documents = SimpleDirectoryReader(input_files=["you_can_put_any_document.pdf"]).load_data()

# step 6: Split sentences for indexing
# Break the document into smaller chunks for efficient processing.

splitter = SentenceSplitter(chunk_size=2048)
nodes = splitter.get_nodes_from_documents(documents)

# step 7: Create the summary index
# Generate an index to store and manage summaries.

summary_index = SummaryIndex(nodes)

# step 8: Configure the query engine
# Set up the query engine for generating tree-based summaries.

summary_query_engine = summary_index.as_query_engine(
    response_mode="tree_summarize",
    use_async=True
)

# step 9: Generate the summary
# Query the engine to generate a detailed summary of the document.

response = summary_query_engine.query("Give a detailed summary of the given document.")

# step 10: Extract the summary text
# Retrieve the text content from the query response.

summary_text = response.content if hasattr(response, 'content') else str(response)

# step 11: Persist the summary index
# Save the summary index for future use.

summary_index.storage_context.persist("summary_index")

# step 12: Save the summary to a Word document
# Write the summary into a Word document with a heading.

doc = Document()
doc.add_heading("Summary of minutes of meeting document", level=1)
doc.add_paragraph(summary_text)  # Add the generated summary as content.
doc.save("summary.docx")
