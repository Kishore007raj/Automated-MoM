import os
from flask import Flask, render_template, request, send_from_directory
from Audio_To_MoM import transcribe_audio, create_docx_from_text
from docx import Document
from llama_index.core import SummaryIndex, Settings
from llama_index.core.node_parser import SentenceSplitter
from llama_index.core import SimpleDirectoryReader
import warnings

# Suppress warnings
warnings.filterwarnings("ignore")

app = Flask(__name__)

# Ensure the output directory exists
OUTPUT_DIR = "outputs"
if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)


# Home route that renders the file upload form
@app.route('/')
def home():
    # Clear any previously rendered download link by not passing `docx_path`
    return render_template('index.html')


# Route to handle file upload and process the audio
@app.route('/generate_minutes', methods=['POST'])
def generate_minutes():
    if 'audio_file' not in request.files:
        return render_template('index.html', error="No file part")

    file = request.files['audio_file']

    if file.filename == '':
        return render_template('index.html', error="No selected file")

    if file:
        file_path = os.path.join("uploads", file.filename)
        file.save(file_path)

        # Perform transcription
        print("Transcribing audio...")
        transcribed_text = transcribe_audio(file_path)

        # Create the transcription document
        transcription_docx_path = os.path.join(OUTPUT_DIR, "transcription.docx")
        create_docx_from_text(transcribed_text, transcription_docx_path)

        # Perform summarization
        print("Generating summary...")
        documents = SimpleDirectoryReader(input_files=[transcription_docx_path]).load_data()
        splitter = SentenceSplitter(chunk_size=2048)
        nodes = splitter.get_nodes_from_documents(documents)

        summary_index = SummaryIndex(nodes)
        summary_query_engine = summary_index.as_query_engine(
            response_mode="tree_summarize",
            use_async=True
        )

        response = summary_query_engine.query("Give a detailed summary of the given document.")
        summary_text = response.content if hasattr(response, 'content') else str(response)

        # Create the summary document
        summary_docx_path = os.path.join(OUTPUT_DIR, "MoM.docx")
        doc = Document()
        doc.add_heading("Minutes of Meeting", 0)
        doc.add_paragraph(summary_text)  # Add the generated summary as content
        doc.save(summary_docx_path)

        # Return the download link for the new file
        return render_template('index.html', docx_path="MoM.docx")


# Route to handle file download
@app.route('/download/<filename>')
def download_file(filename):
    return send_from_directory(OUTPUT_DIR, filename)


if __name__ == '__main__':
    app.run(debug=True)
